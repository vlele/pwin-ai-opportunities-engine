from __future__ import annotations

import io
import json
import mimetypes
import os
from collections import Counter
from pathlib import Path
import re
from typing import Any
import urllib.error
import urllib.parse
import urllib.request
import zipfile
from xml.etree import ElementTree as ET

try:
    from PyPDF2 import PdfReader
except ImportError:  # pragma: no cover - optional dependency
    PdfReader = None

try:
    import fitz
except ImportError:  # pragma: no cover - optional dependency
    fitz = None

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency
    Document = None
from common.runtime import USER_AGENT


SEARCH_URL = "https://api.sam.gov/opportunities/v2/search"
PUBLIC_RESOURCES_URL = "https://sam.gov/api/prod/opps/v3/opportunities/{notice_id}/resources"
PUBLIC_RESOURCE_DOWNLOAD_URL = "https://sam.gov/api/prod/opps/v3/opportunities/resources/files/{resource_id}/download"
CATEGORY_PRIORITY = {
    "statement_of_work": 0,
    "solicitation": 1,
    "instructions_evaluation": 2,
    "questions_answers": 3,
    "amendment": 4,
    "schedule": 5,
    "pricing": 6,
    "subcontracting": 7,
    "other": 8,
}
SNIPPET_KEYWORDS = (
    "scope",
    "objective",
    "requirements",
    "deliver",
    "deliverables",
    "evaluation",
    "period of performance",
    "transition",
    "incumbent",
    "current contractor",
)
STRUCTURED_LINE_MARKERS = (
    "clin ",
    "subclin",
    "task ",
    "subtask",
    "pws",
    "performance work statement",
    "statement of objectives",
    "statement of work",
    "deliverable",
    "transition",
    "period of performance",
    "shall",
    "must",
    "provide",
    "maintain",
    "report",
    "staff",
)
TOC_LINE_RE = re.compile(r"\.{5,}\s*\d+\s*$")
SECTION_HEADING_RE = re.compile(r"^\s*\d+(?:\.\d+){0,3}\s+[A-Z][A-Za-z0-9/\-() ,]+$")
PREFERRED_SECTION_HINTS: tuple[tuple[str, int], ...] = (
    ("statement of work", 6),
    ("performance work statement", 6),
    ("description/specifications/work statement", 6),
    ("description specifications work statement", 6),
    ("statement of objectives", 5),
    ("statement of objective", 5),
    ("scope", 5),
    ("general requirements", 5),
    ("requirements", 4),
    ("specific tasks", 5),
    ("tasks", 4),
    ("deliverables", 5),
    ("period of performance", 4),
    ("background", 3),
    ("purpose", 3),
    ("evaluation factors", 3),
    ("instructions to offerors", 3),
    ("instructions to offeror", 3),
)
ATTACHMENT_TABLE_ROW_RE = re.compile(
    r"^\s*(?:CLIN|SubCLIN|Task|Subtask|AQL|PRS|Performance Objective|Deliverable|Transition|SLIN)\b",
    re.IGNORECASE,
)
TABLE_CELL_SPLIT_RE = re.compile(r"\s*(?:\||\t| {2,})\s*")
ATTACHMENT_MATRIX_MARKERS = (
    "clin",
    "subclin",
    "task",
    "subtask",
    "matrix",
    "deliverable",
    "aql",
    "prs",
    "acceptable quality level",
    "performance requirement",
)
ATTACHMENT_ACCEPTANCE_MARKERS = (
    "acceptance criteria",
    "acceptable quality level",
    "aql",
    "inspection",
    "surveillance method",
    "performance threshold",
    "quality level",
)
ATTACHMENT_INCENTIVE_MARKERS = (
    "incentive",
    "disincentive",
    "remedy",
    "deduction",
    "service credit",
    "liquidated damages",
)
ATTACHMENT_PRICING_MARKERS = (
    "unit price",
    "extended price",
    "total evaluated price",
    "price schedule",
    "rate card",
    "hourly rate",
    "fully burdened",
    "labor rate",
    "cost proposal",
    "price proposal",
    "contract line item",
)
SECTION_BLOCK_HINTS = (
    "statement of objectives",
    "statement of objective",
    "statement of work",
    "performance work statement",
    "description/specifications/work statement",
    "description specifications work statement",
    "scope",
    "general requirements",
    "requirements",
    "deliverables",
    "specific tasks",
    "tasks",
    "task order",
    "period of performance",
    "background",
    "purpose",
    "evaluation factors",
    "instructions to offerors",
    "instructions to offeror",
)
SECTION_BLOCK_SKIP_HINTS = (
    "table of contents",
    "clauses incorporated by reference",
    "representations and certifications",
)
SECTION_HEADING_HINT_RE = re.compile(
    r"^\s*(?:section\s+[a-z0-9]+(?:\s*[-–]\s*)?|\d+(?:\.\d+){0,3}\s+)[a-z].{2,140}$",
    re.IGNORECASE,
)
ACTION_VERB_HINTS = (
    "shall",
    "must",
    "provide",
    "deliver",
    "maintain",
    "support",
    "perform",
    "report",
    "transition",
)
ROOT_SCOPE_HEADINGS = (
    "statement of work",
    "performance work statement",
    "description/specifications/work statement",
    "description specifications work statement",
    "statement of objectives",
    "statement of objective",
)
COMPACT_SECTION_PREFERRED_MARKERS = (
    "shall",
    "must",
    "provide",
    "deliver",
    "maintain",
    "perform",
    "support",
    "report",
    "transition",
    "inspect",
    "review",
    "manage",
    "task",
    "deliverable",
    "clin",
    "subclin",
    "scope",
    "requirement",
    "acceptance",
    "quality",
    "aql",
    "prs",
)
INLINE_HEADING_BODY_RE = re.compile(
    r"^(\d+(?:\.\d+){0,3}\s+(?:Scope|General Requirements|Deliverables?|Tasks?|Specific Tasks?|Performance Objectives?|Requirements?))\s+(?=(?:The contractor|Contractor|Provide|Maintain|Deliver|Perform|Support|Report|Inspect|Review|Manage)\b)",
    re.IGNORECASE,
)


def _normalize_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _dedupe_strings(values: list[str]) -> list[str]:
    seen: set[str] = set()
    deduped: list[str] = []
    for value in values:
        cleaned = _normalize_text(value)
        if not cleaned:
            continue
        key = cleaned.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(cleaned)
    return deduped


def _normalize_preserve_lines(value: object, max_chars: int = 8000) -> str:
    lines: list[str] = []
    for raw_line in str(value or "").replace("\r", "\n").split("\n"):
        cleaned = re.sub(r"[ \t]+", " ", raw_line).strip()
        if cleaned:
            lines.append(cleaned)
    return "\n".join(lines)[:max_chars]


def _is_toc_like_line(line: str) -> bool:
    lowered = line.lower()
    return "table of contents" in lowered or TOC_LINE_RE.search(line) is not None or line.count(".") >= 12


def _page_is_toc_like(text: str) -> bool:
    lines = [line.strip() for line in str(text or "").splitlines() if line.strip()]
    if not lines:
        return False
    toc_like = sum(1 for line in lines if _is_toc_like_line(line))
    return any("table of contents" in line.lower() for line in lines) or (toc_like >= 4 and toc_like >= max(2, len(lines) // 3))


def _section_hint_weight(line: str) -> int:
    lowered = line.lower()
    if any(skip in lowered for skip in SECTION_BLOCK_SKIP_HINTS):
        return 0
    return max((weight for hint, weight in PREFERRED_SECTION_HINTS if hint in lowered), default=0)


def _normalized_attachment_name(filename: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", str(filename or "").lower()).strip()


def _preferred_pdf_section_excerpt(text: str, max_chars: int = 12000) -> str:
    lines = [line.strip() for line in _normalize_preserve_lines(text, max_chars=30000).splitlines() if line.strip()]
    if not lines:
        return ""
    heading_indexes = [
        index
        for index, line in enumerate(lines)
        if (SECTION_HEADING_RE.match(line) or SECTION_HEADING_HINT_RE.match(line)) and not _is_toc_like_line(line)
    ]
    if not heading_indexes:
        return ""
    sections: list[tuple[int, int, str]] = []
    for position, start_index in enumerate(heading_indexes):
        heading = lines[start_index]
        hint_weight = _section_hint_weight(heading)
        if hint_weight <= 0:
            continue
        next_heading_index = heading_indexes[position + 1] if position + 1 < len(heading_indexes) else len(lines)
        captured: list[str] = [heading]
        action_hits = 0
        structured_hits = 0
        for line in lines[start_index + 1:next_heading_index]:
            lower = line.lower()
            if _is_toc_like_line(line):
                continue
            captured.append(line)
            if any(token in lower for token in ACTION_VERB_HINTS):
                action_hits += 1
            if any(token in lower for token in STRUCTURED_LINE_MARKERS):
                structured_hits += 1
            if len(" ".join(captured)) >= 2200:
                break
        block = "\n".join(captured).strip()
        if len(block) < 80:
            continue
        score = (hint_weight * 10) + min(action_hits, 8) + min(structured_hits, 8)
        sections.append((score, start_index, block))
    if not sections:
        return ""
    selected: list[str] = []
    seen: set[str] = set()
    total_chars = 0
    for _, _, block in sorted(sections, key=lambda item: (-item[0], item[1])):
        key = _normalize_text(block[:240])
        if not key or key in seen:
            continue
        seen.add(key)
        selected.append(block)
        total_chars += len(block) + 2
        if len(selected) >= 6 or total_chars >= max_chars:
            break
    return "\n\n".join(selected)[:max_chars]


def _filename_from_headers(url: str, headers: Any) -> str:
    disposition = headers.get("Content-Disposition", "") if headers else ""
    match = re.search(r"filename\*?=(?:UTF-8''|)([^;]+)", disposition, re.IGNORECASE)
    if match:
        filename = urllib.parse.unquote_plus(match.group(1).strip().strip('"'))
        if filename:
            return filename
    path_name = urllib.parse.urlparse(url).path.rsplit("/", 1)[-1]
    return urllib.parse.unquote_plus(path_name or "attachment.bin")


def _attachment_category(filename: str) -> str:
    lowered = _normalized_attachment_name(filename)
    if (
        " sow " in f" {lowered} "
        or "statement of work" in lowered
        or "performance work statement" in lowered
        or "draft pws" in lowered
        or " pws " in f" {lowered} "
        or "soo" in lowered
        or "statement of objectives" in lowered
        or "statement of objective" in lowered
    ):
        return "statement_of_work"
    if "rftop" in lowered or "evaluation" in lowered or "instruction" in lowered or "section l" in lowered or "section m" in lowered:
        return "instructions_evaluation"
    if "question" in lowered or "q&a" in lowered or "qanda" in lowered:
        return "questions_answers"
    if "amendment" in lowered:
        return "amendment"
    if "schedule" in lowered:
        return "schedule"
    if "price" in lowered or "pricing" in lowered or "cost" in lowered:
        return "pricing"
    if "subcontract" in lowered:
        return "subcontracting"
    if ("rfp" in lowered or "rfq" in lowered or "solicitation" in lowered or "soliciation" in lowered):
        return "solicitation"
    return "other"


def _refine_attachment_category(initial_category: str, filename: str, text: str) -> str:
    lower = f"{_normalized_attachment_name(filename)} {_normalize_text(text[:5000]).lower()}".strip()
    if not lower:
        return initial_category
    if any(token in lower for token in ("statement of work", "performance work statement", "statement of objectives", "statement of objective", " draft pws ", " section c - description/specifications/work statement ")):
        return "statement_of_work"
    if any(token in lower for token in ("section l", "section m", "evaluation factors for award", "instructions to offerors", "instruction to offerors", "best value", "past performance questionnaire")):
        return "instructions_evaluation"
    if any(token in lower for token in ("question matrix", "questions and answers", "question responses", "projnet question responses", "q and a", "offeror questions")):
        return "questions_answers"
    if any(token in lower for token in ("amendment of solicitation", "amendment ", "sf30", "modification of contract")):
        return "amendment"
    if initial_category == "other" and any(token in lower for token in ("solicitation/award", "request for quotation", "request for proposal", "combined synopsis", "combined synopsis/solicitation")):
        return "solicitation"
    return initial_category


def _extract_pdf_text_pypdf2(data: bytes, max_pages: int = 40) -> str:
    if PdfReader is None:
        raise RuntimeError("PyPDF2 unavailable")
    reader = PdfReader(io.BytesIO(data))
    page_texts: list[str] = []
    for page in reader.pages[:max_pages]:
        page_texts.append(page.extract_text() or "")
    filtered_pages = [page for page in page_texts if not _page_is_toc_like(page)]
    return "\n".join(filtered_pages or page_texts)


def _extract_pdf_text_fitz(data: bytes, max_pages: int = 40) -> str:
    if fitz is None:
        raise RuntimeError("fitz unavailable")
    document = fitz.open(stream=data, filetype="pdf")
    page_texts: list[str] = []
    for page_index in range(min(len(document), max_pages)):
        page = document.load_page(page_index)
        page_texts.append(page.get_text("text") or "")
    filtered_pages = [page for page in page_texts if not _page_is_toc_like(page)]
    return "\n".join(filtered_pages or page_texts)


def _pdf_text_quality_score(text: str) -> tuple[int, int, int]:
    cleaned = _normalize_preserve_lines(text, max_chars=60000)
    lines = [line for line in cleaned.splitlines() if line.strip()]
    heading_count = sum(1 for line in lines if SECTION_HEADING_RE.match(line) or SECTION_HEADING_HINT_RE.match(line))
    action_count = sum(1 for line in lines if any(term in line.lower() for term in ACTION_VERB_HINTS))
    return (len(cleaned), heading_count, action_count)


def _extract_pdf_text(data: bytes, max_pages: int = 40) -> str:
    candidates: list[str] = []
    errors: list[Exception] = []
    for extractor in (_extract_pdf_text_pypdf2, _extract_pdf_text_fitz):
        try:
            extracted = extractor(data, max_pages=max_pages)
        except Exception as exc:  # pragma: no cover - defensive fallback
            errors.append(exc)
            continue
        if extracted.strip():
            candidates.append(extracted)
    if not candidates:
        if errors:
            raise errors[0]
        raise RuntimeError("No PDF extractor available")
    combined = max(candidates, key=_pdf_text_quality_score)
    preferred = _preferred_pdf_section_excerpt(combined, max_chars=12000)
    if preferred:
        return f"{preferred}\n\n{combined}"
    return combined


def _extract_docx_text(data: bytes) -> str:
    if Document is None:
        raise RuntimeError("python-docx unavailable")
    document = Document(io.BytesIO(data))
    values = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [_normalize_text(cell.text) for cell in row.cells if _normalize_text(cell.text)]
            if not cells:
                continue
            values.append(" | ".join(cells[:10]))
    return "\n".join(values)


def _extract_xlsx_text(data: bytes, max_strings: int = 120) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        names = archive.namelist()
        sheet_names: list[str] = []
        shared_strings: list[str] = []
        workbook_rels: dict[str, str] = {}
        if "xl/workbook.xml" in names:
            workbook = ET.fromstring(archive.read("xl/workbook.xml"))
            ns = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
            sheet_names = [node.attrib.get("name", "") for node in workbook.findall(".//main:sheets/main:sheet", ns)]
            rel_ns = {"rel": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
            if "xl/_rels/workbook.xml.rels" in names:
                rel_root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
                workbook_rels = {
                    node.attrib.get("Id", ""): node.attrib.get("Target", "")
                    for node in rel_root.findall(".//rel:Relationship", rel_ns)
                }
        if "xl/sharedStrings.xml" in names:
            strings_root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
            for node in strings_root.iter():
                if node.tag.endswith("}t") and node.text:
                    shared_strings.append(node.text.strip())
        lines = []
        if sheet_names:
            lines.append(f"Workbook sheets: {', '.join(sheet_names[:8])}")
        if shared_strings:
            sample = [value for value in shared_strings if value][:max_strings]
            lines.append("Shared strings: " + " | ".join(sample[:40]))
        if "xl/workbook.xml" in names:
            workbook = ET.fromstring(archive.read("xl/workbook.xml"))
            ns = {
                "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
                "rel": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            }
            sheet_nodes = workbook.findall(".//main:sheets/main:sheet", ns)
            for sheet_node in sheet_nodes[:4]:
                sheet_name = sheet_node.attrib.get("name", "") or "Sheet"
                rel_id = sheet_node.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", "")
                target = workbook_rels.get(rel_id, "")
                if not target:
                    continue
                sheet_path = f"xl/{target.lstrip('./')}"
                if sheet_path not in names:
                    continue
                root = ET.fromstring(archive.read(sheet_path))
                row_texts: list[str] = []
                for row in root.findall(".//main:sheetData/main:row", ns)[:20]:
                    values: list[str] = []
                    for cell in row.findall("main:c", ns):
                        raw_value = ""
                        cell_type = cell.attrib.get("t", "")
                        if cell_type == "inlineStr":
                            raw_value = "".join(node.text or "" for node in cell.findall(".//main:t", ns)).strip()
                        else:
                            value_node = cell.find("main:v", ns)
                            raw_value = (value_node.text or "").strip() if value_node is not None else ""
                            if raw_value and cell_type == "s":
                                try:
                                    raw_value = shared_strings[int(raw_value)]
                                except (ValueError, IndexError):
                                    pass
                        cleaned = _normalize_text(raw_value)
                        if cleaned:
                            values.append(cleaned)
                    if values:
                        row_texts.append(" | ".join(values[:10]))
                    if len(row_texts) >= 6:
                        break
                if row_texts:
                    lines.append(f"{sheet_name}: " + " || ".join(row_texts[:4]))
        return "\n".join(lines)


def _extract_plain_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def _extract_attachment_text(filename: str, content_type: str, data: bytes) -> tuple[str, str]:
    lowered = filename.lower()
    if lowered.endswith(".pdf") and PdfReader is None:
        return "", "dependency_missing:PyPDF2"
    if lowered.endswith(".docx") and Document is None:
        return "", "dependency_missing:python-docx"
    try:
        if lowered.endswith(".pdf"):
            return _extract_pdf_text(data), "parsed_pdf"
        if lowered.endswith(".docx"):
            return _extract_docx_text(data), "parsed_docx"
        if lowered.endswith(".xlsx"):
            return _extract_xlsx_text(data), "parsed_xlsx"
        if lowered.endswith((".txt", ".csv", ".json", ".xml", ".md")):
            return _extract_plain_text(data), "parsed_text"
        if "text" in content_type:
            return _extract_plain_text(data), "parsed_text"
    except Exception as exc:  # pragma: no cover - defensive fallback
        return "", f"parse_error:{exc.__class__.__name__}"
    return "", "unsupported"


def _pdf_page_count(data: bytes) -> int:
    if PdfReader is None:
        if fitz is None:
            return 0
    try:
        if PdfReader is not None:
            return len(PdfReader(io.BytesIO(data)).pages)
    except Exception:
        pass
    if fitz is None:
        return 0
    try:
        return len(fitz.open(stream=data, filetype="pdf"))
    except Exception:  # pragma: no cover - defensive fallback
        return 0


def _is_structure_heading(line: str) -> bool:
    cleaned = _normalize_text(line)
    if not cleaned or _is_toc_like_line(cleaned):
        return False
    lower = cleaned.lower()
    if any(hint == lower or lower.startswith(f"{hint} ") for hint in SECTION_BLOCK_HINTS):
        return True
    if SECTION_HEADING_RE.match(cleaned) or SECTION_HEADING_HINT_RE.match(cleaned):
        return True
    return False


def _heading_title(line: str) -> str:
    cleaned = _normalize_text(line)
    match = re.match(r"^(?:section\s+[a-z0-9]+(?:\s*[-–]\s*)?|\d+(?:\.\d+){0,3}\s+)?(.+)$", cleaned, re.IGNORECASE)
    title = match.group(1) if match else cleaned
    return title.strip(" .:-")


def _split_heading_and_inline_body(line: str) -> tuple[str, str]:
    cleaned = _normalize_text(line)
    if not cleaned:
        return "", ""
    cleaned = INLINE_HEADING_BODY_RE.sub(r"\1; ", cleaned)
    lower = cleaned.lower()
    for hint in sorted(SECTION_BLOCK_HINTS, key=len, reverse=True):
        if lower == hint:
            return cleaned, ""
        for separator in (":", ";", " - ", ". "):
            prefix = f"{hint}{separator}"
            if lower.startswith(prefix):
                title = cleaned[: len(hint)]
                inline_body = cleaned[len(prefix):].strip(" ;:-")
                return title, inline_body
        if lower.startswith(f"{hint} "):
            remainder = cleaned[len(hint):].strip(" ;:-")
            if remainder and any(marker in remainder.lower() for marker in ACTION_VERB_HINTS):
                return cleaned[: len(hint)], remainder
    return _heading_title(cleaned), ""


def _section_block_priority(title: str, body: str) -> int:
    lower = f"{title} {body}".lower()
    score = 0
    for hint in SECTION_BLOCK_HINTS:
        if hint in lower:
            score += 4
    for verb in ACTION_VERB_HINTS:
        if verb in lower:
            score += 1
    if any(marker in lower for marker in ("clin", "task", "deliverable", "transition", "period of performance", "base period", "option year")):
        score += 3
    return score


def _table_row_kind(text: str) -> str:
    lower = str(text or "").lower()
    if lower.startswith(("clin ", "subclin", "slin")):
        return "clin"
    if lower.startswith(("task ", "subtask", "performance objective")):
        return "task"
    if any(marker in lower for marker in ATTACHMENT_PRICING_MARKERS):
        return "pricing"
    if any(marker in lower for marker in ATTACHMENT_ACCEPTANCE_MARKERS):
        return "acceptance"
    if any(marker in lower for marker in ATTACHMENT_INCENTIVE_MARKERS):
        return "remedy"
    if any(marker in lower for marker in ATTACHMENT_MATRIX_MARKERS):
        return "matrix"
    return "table"


def _matrix_row_record(raw_line: str, line_index: int) -> dict[str, Any] | None:
    normalized_row = _normalize_text(raw_line)
    normalized_row = TABLE_CELL_SPLIT_RE.sub("; ", normalized_row).strip(" ;:-")
    if len(normalized_row) < 24:
        return None
    cells = [cell.strip(" ;:-") for cell in TABLE_CELL_SPLIT_RE.split(_normalize_text(raw_line)) if cell.strip(" ;:-")]
    return {
        "line_index": line_index,
        "kind": _table_row_kind(normalized_row),
        "label": cells[0] if cells else normalized_row[:80],
        "text": normalized_row[:420],
        "cells": cells[:10],
    }


def _table_blocks(row_records: list[dict[str, Any]], max_blocks: int = 8) -> list[dict[str, Any]]:
    if not row_records:
        return []
    blocks: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for row in sorted(row_records, key=lambda item: int(item.get("line_index", 0) or 0)):
        kind = str(row.get("kind", "table") or "table")
        line_index = int(row.get("line_index", 0) or 0)
        if (
            current is None
            or kind != str(current.get("kind", "table") or "table")
            or line_index - int(current.get("last_line_index", line_index) or line_index) > 2
        ):
            current = {
                "kind": kind,
                "title": f"{kind.replace('_', ' ').title()} table",
                "rows": [],
                "last_line_index": line_index,
            }
            blocks.append(current)
        current["rows"].append(str(row.get("text", "") or "").strip())
        current["last_line_index"] = line_index
    rendered: list[dict[str, Any]] = []
    for block in blocks[:max_blocks]:
        rows = [row for row in (block.get("rows", []) or []) if str(row or "").strip()]
        if not rows:
            continue
        rendered.append(
            {
                "kind": str(block.get("kind", "table") or "table"),
                "title": str(block.get("title", "Table") or "Table"),
                "row_count": len(rows),
                "rows": rows[:10],
            }
        )
    return rendered


def _section_graph(raw_lines: list[str], max_nodes: int = 24) -> list[dict[str, Any]]:
    nodes: list[dict[str, Any]] = []
    stack: list[dict[str, Any]] = []
    for raw_line in raw_lines:
        if not _is_structure_heading(raw_line):
            continue
        title, _ = _split_heading_and_inline_body(raw_line)
        cleaned_title = _heading_title(title).strip()
        if not cleaned_title:
            continue
        number_match = re.match(r"^\s*(\d+(?:\.\d+){0,3})\b", str(raw_line or "").strip())
        level = number_match.group(1).count(".") + 1 if number_match else 1
        while stack and int(stack[-1].get("level", 1) or 1) >= level:
            stack.pop()
        parent_title = str(stack[-1].get("title", "") or "").strip() if stack else ""
        node = {
            "title": cleaned_title[:180],
            "level": level,
            "parent_title": parent_title,
        }
        nodes.append(node)
        stack.append(node)
        if len(nodes) >= max_nodes:
            break
    return nodes


def _structure_parse_warnings(text: str, row_records: list[dict[str, Any]], section_blocks: list[dict[str, str]]) -> list[str]:
    lower = str(text or "").lower()
    warnings: list[str] = []
    has_pricing_markers = any(marker in lower for marker in ATTACHMENT_PRICING_MARKERS)
    has_acceptance_markers = any(marker in lower for marker in ATTACHMENT_ACCEPTANCE_MARKERS)
    has_remedy_markers = any(marker in lower for marker in ATTACHMENT_INCENTIVE_MARKERS)
    has_clin_markers = any(marker in lower for marker in ("clin ", "subclin", "slin"))
    kinds = Counter(str(row.get("kind", "") or "") for row in row_records)
    if has_pricing_markers and kinds.get("pricing", 0) == 0:
        warnings.append("pricing_markers_without_structured_rows")
    if has_acceptance_markers and kinds.get("acceptance", 0) == 0:
        warnings.append("acceptance_markers_without_structured_rows")
    if has_remedy_markers and kinds.get("remedy", 0) == 0:
        warnings.append("remedy_markers_without_structured_rows")
    if has_clin_markers and kinds.get("clin", 0) == 0:
        warnings.append("clin_markers_without_structured_rows")
    if any("scope" in str(block.get("title", "") or "").lower() for block in section_blocks) and not any(
        "shall" in str(block.get("source_text", "") or "").lower() for block in section_blocks
    ):
        warnings.append("scope_sections_without_action_sentences")
    return _dedupe_strings(warnings)


def _compact_section_body(body_lines: list[str], max_chars: int = 680) -> str:
    candidates: list[tuple[int, str]] = []
    for raw_line in body_lines:
        cleaned = _normalize_text(raw_line)
        if not cleaned or _is_toc_like_line(cleaned):
            continue
        lower = cleaned.lower()
        if lower in ROOT_SCOPE_HEADINGS:
            continue
        if ATTACHMENT_TABLE_ROW_RE.match(cleaned):
            normalized_row = re.sub(r"\s*(?:\||\t| {2,})\s*", "; ", cleaned).strip(" ;:-")
            candidates.append((8, normalized_row))
            continue
        for part in re.split(r"(?<=[.;])\s+|(?<=\))\s+(?=[A-Z])", cleaned):
            sentence = _normalize_text(part).strip(" ;:-")
            if len(sentence) < 24:
                continue
            lower_sentence = sentence.lower()
            score = 0
            if any(marker in lower_sentence for marker in COMPACT_SECTION_PREFERRED_MARKERS):
                score += 4
            if any(marker in lower_sentence for marker in ATTACHMENT_MATRIX_MARKERS):
                score += 2
            if any(marker in lower_sentence for marker in ATTACHMENT_ACCEPTANCE_MARKERS):
                score += 2
            if any(marker in lower_sentence for marker in ATTACHMENT_PRICING_MARKERS):
                score += 2
            if any(marker in lower_sentence for marker in ("mission", "purpose", "background")) and not any(
                verb in lower_sentence for verb in ACTION_VERB_HINTS
            ):
                score -= 3
            if score <= 0:
                continue
            candidates.append((score, sentence))

    if not candidates:
        fallback = _normalize_text(" ".join(body_lines))
        return fallback[:max_chars]

    selected: list[str] = []
    seen: set[str] = set()
    for _, sentence in sorted(candidates, key=lambda item: (-item[0], len(item[1]), item[1])):
        key = sentence.lower()
        if key in seen:
            continue
        if any(key in existing.lower() or existing.lower() in key for existing in selected):
            continue
        seen.add(key)
        selected.append(sentence)
        if len("; ".join(selected)) >= max_chars or len(selected) >= 3:
            break
    return "; ".join(selected)[:max_chars]


def _extract_text_structures(text: str, *, max_sections: int = 12, max_rows: int = 36) -> dict[str, Any]:
    raw_lines = [line.strip() for line in str(text or "").replace("\r", "\n").split("\n") if line.strip()]
    headings: list[str] = []
    section_blocks: list[dict[str, str]] = []
    structured_rows: list[str] = []
    structured_row_records: list[dict[str, Any]] = []
    seen_sections: set[str] = set()
    seen_rows: set[str] = set()

    for line_index, raw_line in enumerate(raw_lines):
        if (
            ATTACHMENT_TABLE_ROW_RE.match(raw_line)
            or "|" in raw_line
            or "\t" in raw_line
            or re.search(r" {2,}", raw_line)
        ):
            row_record = _matrix_row_record(raw_line, line_index)
            if row_record is not None:
                normalized_row = str(row_record.get("text", "") or "")
                row_key = normalized_row.lower()
                if row_key not in seen_rows:
                    seen_rows.add(row_key)
                    structured_rows.append(normalized_row[:420])
                    structured_row_records.append(row_record)

    index = 0
    while index < len(raw_lines):
        line = raw_lines[index]
        if not _is_structure_heading(line):
            index += 1
            continue
        title, inline_body = _split_heading_and_inline_body(line)
        lower_title = title.lower()
        headings.append(title)
        if any(skip in lower_title for skip in SECTION_BLOCK_SKIP_HINTS):
            index += 1
            continue
        body_lines: list[str] = [inline_body] if inline_body else []
        lookahead = index + 1
        while lookahead < len(raw_lines):
            candidate = raw_lines[lookahead]
            if _is_structure_heading(candidate):
                break
            if not _is_toc_like_line(candidate):
                body_lines.append(candidate)
            if len(" ".join(body_lines)) >= 2200:
                break
            lookahead += 1
        body = _normalize_text(" ".join(body_lines))
        compact_body = _compact_section_body(body_lines)
        if lower_title in ROOT_SCOPE_HEADINGS and not compact_body:
            index = lookahead if lookahead > index else index + 1
            continue
        if body and _section_block_priority(title, body) >= 4:
            section_text = _normalize_text(f"{title}; {compact_body or body}")[:1200]
            key = section_text.lower()
            if key not in seen_sections:
                seen_sections.add(key)
                section_blocks.append(
                    {
                        "title": title[:180],
                        "text": section_text,
                        "source_text": _normalize_text(f"{title}; {body}")[:2200],
                    }
                )
        index = lookahead if lookahead > index else index + 1

    section_blocks.sort(key=lambda item: _section_block_priority(item.get("title", ""), item.get("text", "")), reverse=True)
    table_blocks = _table_blocks(structured_row_records)
    matrix_rows = [row for row in structured_row_records if str(row.get("kind", "") or "") in {"clin", "task", "matrix"}]
    pricing_rows = [row for row in structured_row_records if str(row.get("kind", "") or "") == "pricing"]
    acceptance_rows = [row for row in structured_row_records if str(row.get("kind", "") or "") == "acceptance"]
    remedy_rows = [row for row in structured_row_records if str(row.get("kind", "") or "") == "remedy"]
    parse_warnings = _structure_parse_warnings(text, structured_row_records, section_blocks)
    return {
        "headings": _dedupe_strings(headings)[:24],
        "section_blocks": section_blocks[:max_sections],
        "structured_rows": structured_rows[:max_rows],
        "table_blocks": table_blocks,
        "matrix_rows": matrix_rows[:max_rows],
        "pricing_rows": pricing_rows[:max_rows],
        "acceptance_rows": acceptance_rows[:max_rows],
        "remedy_rows": remedy_rows[:max_rows],
        "section_graph": _section_graph(raw_lines),
        "parse_warnings": parse_warnings,
    }


def _attachment_analysis(
    *,
    filename: str,
    content_type: str,
    category: str,
    parser_status: str,
    text: str,
    data: bytes,
) -> dict[str, Any]:
    lines = [line.strip() for line in _normalize_preserve_lines(text, max_chars=40000).splitlines() if line.strip()]
    lower_lines = [line.lower() for line in lines]
    joined_lower = "\n".join(lower_lines)
    char_count = len(text.strip())
    line_count = len(lines)
    page_count = _pdf_page_count(data) if filename.lower().endswith(".pdf") else 0
    table_like_line_count = sum(
        1
        for line in lines
        if len(line) >= 20 and ("|" in line or "\t" in line or re.search(r" {2,}", line) or ATTACHMENT_TABLE_ROW_RE.match(line))
    )
    matrix_marker_count = sum(1 for line in lower_lines if any(marker in line for marker in ATTACHMENT_MATRIX_MARKERS))
    acceptance_marker_count = sum(1 for line in lower_lines if any(marker in line for marker in ATTACHMENT_ACCEPTANCE_MARKERS))
    incentive_marker_count = sum(1 for line in lower_lines if any(marker in line for marker in ATTACHMENT_INCENTIVE_MARKERS))
    pricing_marker_count = sum(1 for line in lower_lines if any(marker in line for marker in ATTACHMENT_PRICING_MARKERS))
    flags: list[str] = []

    if parser_status.startswith(("parse_error", "dependency_missing")) or parser_status == "unsupported":
        flags.append("parse_failed_or_unsupported")
    if filename.lower().endswith(".pdf"):
        if page_count >= 4 and char_count < max(900, page_count * 180):
            flags.append("ocr_or_image_heavy_pdf")
        elif len(data) >= 600_000 and char_count < 1200:
            flags.append("ocr_or_image_heavy_pdf")
    if parser_status.startswith("parsed") and char_count < 400:
        flags.append("thin_text_extraction")
    if table_like_line_count >= 5 or matrix_marker_count >= 4:
        flags.append("table_or_matrix_heavy")
    if sum(1 for line in lines if ATTACHMENT_TABLE_ROW_RE.match(line)) >= 3:
        flags.append("clin_or_task_matrix_visible")
    if acceptance_marker_count >= 2:
        flags.append("acceptance_or_aql_visible")
    if incentive_marker_count >= 2:
        flags.append("incentive_or_remedy_visible")
    if category == "pricing" or filename.lower().endswith(".xlsx") or (pricing_marker_count >= 3 and table_like_line_count >= 2):
        flags.append("pricing_sheet_or_rate_table")
    if category in {"statement_of_work", "solicitation"} and "statement of work" in joined_lower and char_count < 900:
        flags.append("scope_document_underparsed")

    review_required = any(
        flag in flags
        for flag in (
            "parse_failed_or_unsupported",
            "ocr_or_image_heavy_pdf",
            "thin_text_extraction",
            "scope_document_underparsed",
        )
    )
    return {
        "text_char_count": char_count,
        "line_count": line_count,
        "page_count": page_count,
        "table_like_line_count": table_like_line_count,
        "matrix_marker_count": matrix_marker_count,
        "acceptance_marker_count": acceptance_marker_count,
        "incentive_marker_count": incentive_marker_count,
        "pricing_marker_count": pricing_marker_count,
        "analysis_flags": flags,
        "review_required": review_required,
    }


def _attachment_snippets(text: str, max_snippets: int = 6) -> list[str]:
    structured_lines = _normalize_preserve_lines(text, max_chars=12000).splitlines()
    ranked_structured: list[str] = []
    for line in structured_lines:
        lower = line.lower()
        if len(line) < 40:
            continue
        if _is_toc_like_line(line):
            continue
        if any(marker in lower for marker in STRUCTURED_LINE_MARKERS):
            ranked_structured.append(line[:360])
    if ranked_structured:
        return ranked_structured[:max_snippets]

    cleaned = _normalize_text(text)
    if not cleaned:
        return []
    parts = re.split(r"(?<=[.!?])\s+|\n+", cleaned)
    ranked: list[str] = []
    for part in parts:
        snippet = part.strip()
        if len(snippet) < 60:
            continue
        if any(keyword in snippet.lower() for keyword in SNIPPET_KEYWORDS):
            ranked.append(snippet[:320])
    if ranked:
        return ranked[:max_snippets]
    return [part.strip()[:320] for part in parts if part.strip()][:max_snippets]


def _download_attachment(
    url: str,
    *,
    filename_hint: str = "",
    content_type_hint: str = "",
    timeout: int = 45,
    max_bytes: int = 25_000_000,
) -> dict[str, Any]:
    request = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
    with urllib.request.urlopen(request, timeout=timeout) as response:
        data = response.read(max_bytes + 1)
        truncated = len(data) > max_bytes
        if truncated:
            data = data[:max_bytes]
        filename = filename_hint or _filename_from_headers(url, response.headers)
        content_type = content_type_hint or response.headers.get("Content-Type", "application/octet-stream")
    return _attachment_record_from_bytes(
        url=url,
        filename=filename,
        content_type=content_type,
        data=data,
        truncated=truncated,
    )


def _attachment_record_from_bytes(
    *,
    url: str,
    filename: str,
    content_type: str,
    data: bytes,
    truncated: bool,
    local_path: str = "",
) -> dict[str, Any]:
    text, parser_status = _extract_attachment_text(filename, content_type, data)
    structured_excerpt = _normalize_preserve_lines(text, max_chars=120000)
    category = _refine_attachment_category(_attachment_category(filename), filename, structured_excerpt)
    analysis = _attachment_analysis(
        filename=filename,
        content_type=content_type,
        category=category,
        parser_status=parser_status,
        text=structured_excerpt,
        data=data,
    )
    record = {
        "url": url,
        "filename": filename,
        "content_type": content_type,
        "size_bytes": len(data),
        "truncated": truncated,
        "parser_status": parser_status,
        "text_excerpt": structured_excerpt[:18000],
        "structured_text_excerpt": structured_excerpt,
        "category": category,
        **analysis,
    }
    structures = _extract_text_structures(structured_excerpt)
    record["headings"] = structures.get("headings", [])
    record["section_blocks"] = structures.get("section_blocks", [])
    record["structured_rows"] = structures.get("structured_rows", [])
    record["snippets"] = _attachment_snippets(
        "\n".join(
            [
                *(block.get("text", "") for block in record["section_blocks"] if isinstance(block, dict)),
                structured_excerpt,
            ]
        )
    )
    if local_path:
        record["local_path"] = local_path
    return record


def load_local_attachments(
    file_paths: list[str] | None,
    *,
    max_attachments: int = 20,
    max_bytes: int = 25_000_000,
) -> dict[str, Any]:
    attachments: list[dict[str, Any]] = []
    errors: list[str] = []
    normalized_paths: list[Path] = []
    for raw_path in file_paths or []:
        candidate = str(raw_path or "").strip()
        if not candidate:
            continue
        path = Path(candidate).expanduser()
        if not path.exists():
            errors.append(f"{path.as_posix()}: file_not_found")
            continue
        if not path.is_file():
            errors.append(f"{path.as_posix()}: not_a_file")
            continue
        normalized_paths.append(path)

    for path in normalized_paths[:max_attachments]:
        try:
            data = path.read_bytes()
            truncated = len(data) > max_bytes
            if truncated:
                data = data[:max_bytes]
            content_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
            attachments.append(
                _attachment_record_from_bytes(
                    url=path.resolve().as_uri(),
                    filename=path.name,
                    content_type=content_type,
                    data=data,
                    truncated=truncated,
                    local_path=path.resolve().as_posix(),
                )
            )
        except Exception as exc:  # pragma: no cover - defensive fallback
            errors.append(f"{path.as_posix()}: {exc}")

    attachments.sort(key=lambda item: (CATEGORY_PRIORITY.get(item.get("category", "other"), 99), item.get("filename", "")))
    status = "ok" if attachments else ("error" if errors else "empty")
    return {
        "status": status,
        "record": {},
        "point_of_contact": [],
        "attachments": attachments,
        "attachments_expected": bool(normalized_paths),
        "record_lookup_status": "local_files",
        "resource_listing_status": "local_files",
        "seeded_resource_links": False,
        "resource_link_count": len(normalized_paths),
        "errors": errors,
    }


def _record_value(record: dict[str, Any], *keys: str) -> Any:
    for key in keys:
        value = record.get(key)
        if value not in (None, ""):
            return value
    return ""


def _fetch_public_notice_record(notice_id: str, solicitation_number: str = "", timeout: int = 30) -> dict[str, Any]:
    api_key = os.getenv("SAM_API_KEY", "").strip()
    if not api_key:
        return {"status": "missing_api_key", "record": {}}
    params = {"api_key": api_key}
    if notice_id:
        params["noticeid"] = notice_id
    elif solicitation_number:
        params["solnum"] = solicitation_number
    else:
        return {"status": "missing_identifier", "record": {}}
    request = urllib.request.Request(
        f"{SEARCH_URL}?{urllib.parse.urlencode(params)}",
        headers={"User-Agent": USER_AGENT},
        method="GET",
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            payload = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        return {"status": "http_error", "detail": detail[:500], "record": {}}
    except Exception as exc:  # pragma: no cover - defensive fallback
        return {"status": "error", "detail": str(exc), "record": {}}

    rows = payload.get("opportunitiesData", [])
    if not isinstance(rows, list) or not rows:
        return {"status": "empty", "record": {}}
    return {"status": "ok", "record": rows[0]}


def _fetch_public_resource_links(notice_id: str, timeout: int = 30) -> dict[str, Any]:
    if not notice_id:
        return {"status": "missing_identifier", "resources": []}
    request = urllib.request.Request(
        PUBLIC_RESOURCES_URL.format(notice_id=urllib.parse.quote(notice_id)),
        headers={"User-Agent": USER_AGENT},
        method="GET",
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            payload = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        return {"status": "http_error", "detail": detail[:500], "resources": []}
    except Exception as exc:  # pragma: no cover - defensive fallback
        return {"status": "error", "detail": str(exc), "resources": []}

    embedded = payload.get("_embedded", {}) if isinstance(payload, dict) else {}
    attachment_groups = embedded.get("opportunityAttachmentList", []) if isinstance(embedded, dict) else []
    resources: list[dict[str, Any]] = []
    seen_urls: set[str] = set()
    for group in attachment_groups if isinstance(attachment_groups, list) else []:
        attachments = group.get("attachments", []) if isinstance(group, dict) else []
        for item in attachments if isinstance(attachments, list) else []:
            if not isinstance(item, dict):
                continue
            resource_id = str(item.get("resourceId", "") or "").strip()
            if not resource_id:
                continue
            url = PUBLIC_RESOURCE_DOWNLOAD_URL.format(resource_id=urllib.parse.quote(resource_id))
            if url in seen_urls:
                continue
            seen_urls.add(url)
            resources.append(
                {
                    "url": url,
                    "filename": str(item.get("name", "") or "").strip(),
                    "content_type": str(item.get("mimeType", "") or "").strip(),
                    "category": _attachment_category(str(item.get("name", "") or "")),
                    "posted_date": str(item.get("postedDate", "") or "").strip(),
                }
            )
    return {"status": "ok" if resources else "empty", "resources": resources}


def fetch_notice_attachments(
    notice_id: str,
    solicitation_number: str = "",
    resource_links: list[str] | None = None,
    point_of_contact: list[dict[str, Any]] | None = None,
    *,
    max_attachments: int = 20,
) -> dict[str, Any]:
    seeded_resource_links = bool(download_targets := [
        {"url": link.strip(), "filename": "", "content_type": ""}
        for link in (resource_links or [])
        if isinstance(link, str) and link.strip()
    ])
    contacts = [item for item in (point_of_contact or []) if isinstance(item, dict)]
    record: dict[str, Any] = {}
    status = "ok"
    errors: list[str] = []
    record_lookup_status = "skipped"
    resource_listing_status = "skipped"

    if not download_targets:
        public_record_result = _fetch_public_notice_record(notice_id, solicitation_number)
        record = public_record_result.get("record", {})
        status = public_record_result.get("status", "error")
        record_lookup_status = status
        if status == "ok" and isinstance(record, dict):
            contacts = record.get("pointOfContact", []) if isinstance(record.get("pointOfContact"), list) else contacts
            download_targets = [
                {"url": link.strip(), "filename": "", "content_type": ""}
                for link in (record.get("resourceLinks") or [])
                if isinstance(link, str) and link.strip()
            ]
        else:
            errors.append(public_record_result.get("detail", "No public notice record was returned."))

    if not download_targets:
        public_resources_result = _fetch_public_resource_links(notice_id)
        status = public_resources_result.get("status", status)
        resource_listing_status = status
        download_targets = public_resources_result.get("resources", []) if isinstance(public_resources_result.get("resources"), list) else []
        if not download_targets and public_resources_result.get("detail"):
            errors.append(public_resources_result.get("detail"))

    attachments_expected = bool(download_targets)
    attachments: list[dict[str, Any]] = []
    for target in download_targets[:max_attachments]:
        if not isinstance(target, dict):
            continue
        try:
            attachments.append(
                _download_attachment(
                    str(target.get("url", "") or ""),
                    filename_hint=str(target.get("filename", "") or ""),
                    content_type_hint=str(target.get("content_type", "") or ""),
                )
            )
        except Exception as exc:  # pragma: no cover - defensive fallback
            errors.append(f"{target.get('url', 'unknown')}: {exc}")

    attachments.sort(key=lambda item: (CATEGORY_PRIORITY.get(item.get("category", "other"), 99), item.get("filename", "")))
    return {
        "status": "ok" if attachments else ("error" if errors else status),
        "record": record,
        "point_of_contact": contacts,
        "attachments": attachments,
        "attachments_expected": attachments_expected,
        "record_lookup_status": record_lookup_status,
        "resource_listing_status": resource_listing_status,
        "seeded_resource_links": seeded_resource_links,
        "resource_link_count": len(download_targets),
        "errors": errors,
    }
